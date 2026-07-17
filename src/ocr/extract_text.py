"""Extract raw text from contract PDFs and DOCX files.

PDFs are handled in two passes per page:
1. Try direct text extraction (pdfplumber) - fast and exact, works for
   "digital" PDFs that have an embedded text layer.
2. If a page yields fewer than OCR_MIN_CHARS_PER_PAGE characters, treat it
   as a scanned image and fall back to Tesseract OCR (pdf2image renders the
   page to an image, pytesseract reads it).

This mirrors how real contract repositories look in practice: a mix of
natively-generated PDFs and scanned/faxed paper agreements.
"""

import logging
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document

from src.common.config import OCR_DPI, OCR_MIN_CHARS_PER_PAGE

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)


@dataclass
class PageResult:
    page_number: int
    text: str
    method: str  # "digital" or "ocr"

    @property
    def char_count(self) -> int:
        return len(self.text)


class OcrDependencyError(RuntimeError):
    """Raised when Tesseract/poppler are needed but not available on this machine."""


def _ocr_page_image(image) -> str:
    try:
        import pytesseract
    except ImportError as exc:  # pragma: no cover - exercised only without the dep
        raise OcrDependencyError(
            "pytesseract is not installed. Run `pip install -r requirements.txt`."
        ) from exc

    try:
        return pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError as exc:
        raise OcrDependencyError(
            "Tesseract OCR binary not found. Install it "
            "(https://github.com/UB-Mannheim/tesseract/wiki on Windows) and either add it "
            "to PATH or set pytesseract.pytesseract.tesseract_cmd / TESSERACT_CMD."
        ) from exc


def page_needs_ocr(extracted_text: str, min_chars: int = OCR_MIN_CHARS_PER_PAGE) -> bool:
    """A page with too little embedded text is assumed to be a scanned image."""
    return len(extracted_text.strip()) < min_chars


def extract_text_from_pdf(pdf_path: Path) -> list[PageResult]:
    """Extract text page-by-page, OCR'ing any page without a usable text layer."""
    results: list[PageResult] = []
    scanned_page_numbers: list[int] = []

    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not page_needs_ocr(text):
                results.append(PageResult(page_number=i, text=text, method="digital"))
            else:
                results.append(PageResult(page_number=i, text="", method="ocr"))
                scanned_page_numbers.append(i)

    if scanned_page_numbers:
        logger.info(
            "%s: %d page(s) have no text layer, running OCR: %s",
            pdf_path.name,
            len(scanned_page_numbers),
            scanned_page_numbers,
        )
        results = _ocr_missing_pages(pdf_path, results, scanned_page_numbers)

    return results


def _ocr_missing_pages(
    pdf_path: Path, results: list[PageResult], page_numbers: list[int]
) -> list[PageResult]:
    try:
        from pdf2image import convert_from_path
    except ImportError as exc:  # pragma: no cover
        raise OcrDependencyError(
            "pdf2image is not installed. Run `pip install -r requirements.txt`."
        ) from exc

    try:
        images = convert_from_path(
            str(pdf_path),
            dpi=OCR_DPI,
            first_page=min(page_numbers),
            last_page=max(page_numbers),
        )
    except Exception as exc:  # poppler missing raises a generic PDFInfoNotInstalledError
        raise OcrDependencyError(
            "Could not rasterize PDF pages for OCR - poppler may not be installed. "
            "On Windows, install poppler and add its `bin/` folder to PATH: "
            "https://github.com/oschwartz10612/poppler-windows"
        ) from exc

    offset = min(page_numbers)
    by_page = {r.page_number: r for r in results}
    for page_num, image in zip(range(offset, offset + len(images)), images):
        if page_num not in page_numbers:
            continue
        ocr_text = _ocr_page_image(image).strip()
        by_page[page_num] = PageResult(page_number=page_num, text=ocr_text, method="ocr")

    return [by_page[r.page_number] for r in results]


def extract_text_from_docx(docx_path: Path) -> str:
    document = Document(docx_path)
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(p for p in paragraphs if p.strip())


def extract_document(path: Path) -> str:
    """Extract full document text, dispatching on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = extract_text_from_pdf(path)
        return "\n\n".join(p.text for p in pages)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported document type: {suffix} ({path})")
